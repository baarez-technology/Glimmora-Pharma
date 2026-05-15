"""
Build the AI API Testing Manual as a single Word document.

Run from the project root:
    python docs/manual/build_api_testing_manual.py

Output:
    docs/AI-API-Testing-Manual.docx

Screenshots come from docs/test-screenshots/ (already captured via Playwright).
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("python-docx is required. Install with: pip install python-docx\n")
    sys.exit(1)


ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "test-screenshots"
OUT = ROOT / "AI-API-Testing-Manual.docx"


BRAND_AMBER = RGBColor(0x8B, 0x69, 0x14)
BRAND_DARK = RGBColor(0x30, 0x2D, 0x29)
GRAY = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x10, 0xB9, 0x81)
RED = RGBColor(0xDC, 0x26, 0x26)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
BLUE = RGBColor(0x0E, 0xA5, 0xE9)


# ── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_AMBER
    run.font.name = "Calibri"


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BRAND_DARK
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "8B6914")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BRAND_AMBER


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BRAND_DARK


def add_body(doc, text, italic=False, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(11)
    if color is not None:
        run.font.color.rgb = color
    return p


def add_rich_paragraph(doc, segments):
    """segments: list of (text, dict(bold, italic, color, code))"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, style in segments:
        run = p.add_run(text)
        run.bold = bool(style.get("bold"))
        run.italic = bool(style.get("italic"))
        if style.get("color"):
            run.font.color.rgb = style["color"]
        if style.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run("")
    p.runs[0].text = ""
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = BRAND_DARK
    # Light shading on the paragraph
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F2EC")
    p_pr.append(shd)


def add_image(doc, path: Path, caption=None, width_inches=6.2):
    if not path.exists():
        add_body(doc, f"[Missing screenshot: {path.name}]", italic=True, color=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cap.add_run(caption)
        crun.italic = True
        crun.font.size = Pt(9.5)
        crun.font.color.rgb = GRAY


def add_table(doc, header_row, body_rows, col_widths=None):
    table = doc.add_table(rows=1 + len(body_rows), cols=len(header_row))
    table.style = "Light Grid Accent 1"
    # header
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "8B6914")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # body
    for r, row in enumerate(body_rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(r.cells):
                    r.cells[i].width = Inches(w)


def add_step_header(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    a = p.add_run(f"Step {num} — ")
    a.bold = True
    a.font.size = Pt(13)
    a.font.color.rgb = BRAND_AMBER
    b = p.add_run(text)
    b.bold = True
    b.font.size = Pt(13)
    b.font.color.rgb = BRAND_DARK


def add_expected(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    a = p.add_run("Expected: ")
    a.bold = True
    a.font.size = Pt(11)
    a.font.color.rgb = GREEN
    b = p.add_run(text)
    b.font.size = Pt(11)
    b.font.color.rgb = BRAND_DARK


def set_landscape_margins(doc):
    for section in doc.sections:
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)


# ── Build ───────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    set_landscape_margins(doc)

    # ─── Cover ──────────────────────────────────────────────────────────────
    add_title(doc, "Pharma Glimmora")
    add_title(doc, "AI / API Testing Manual")
    add_subtitle(doc, "End-to-end Playwright-verified walkthrough — 31 AI backend endpoints")
    add_body(doc, " ")
    add_body(doc,
        "Branch: devAI    ·    Frontend: localhost:3000    ·    "
        "Backend: pharma-glimmora-ai-backend.onrender.com    ·    Generated: 2026-05-15",
        italic=True, color=GRAY)
    add_body(doc, " ")
    add_body(doc,
        "This document walks through every AI backend endpoint exposed via the same-origin /api/ai-proxy "
        "route. Each step shows the navigation, sample inputs, the screenshot you should see, and the expected "
        "HTTP response. The walkthrough exercises 28 endpoints through the UI and 3 via curl (where no UI "
        "surface exists or where binary audio uploads are needed).")

    add_h2(doc, "Result-symbol legend")
    add_table(doc,
        ["Symbol", "Meaning"],
        [
            ["✅", "Verified through the UI walkthrough"],
            ["🛠️", "Verified via curl through the same proxy (no UI surface or binary input needed)"],
            ["🟡", "Silent-204 collapse — proxy converts 404 → 204 deliberately (\"not started yet\")"],
        ],
        col_widths=[1.0, 7.5])

    add_page_break(doc)

    # ─── §1 Prerequisites ───────────────────────────────────────────────────
    add_h1(doc, "1. Prerequisites")

    add_h2(doc, "1.1  Local environment")
    add_table(doc,
        ["Component", "Setting", "Where"],
        [
            ["Node", "≥ 22", "npm run dev"],
            ["Postgres", "local instance named `pharma`", "DATABASE_URL in .env (gitignored)"],
            ["Prisma", "provider = \"postgresql\"", "prisma/schema.prisma"],
            ["Next.js dev", "http://localhost:3000", "npm run dev"],
            ["AI backend", "Render upstream or local FastAPI :8000", "NEXT_PUBLIC_API_URL / BACKEND_URL"],
        ],
        col_widths=[1.6, 3.6, 3.3])

    add_h2(doc, "1.2  First-time DB setup")
    add_code(doc,
        "psql -U postgres -c 'create database pharma;'\n"
        "npx prisma db push --skip-generate --accept-data-loss\n"
        "npx prisma generate\n"
        "npm run db:seed")

    add_h2(doc, "1.3  Seed accounts")
    add_table(doc,
        ["Account", "Role", "Username / Email", "Password"],
        [
            ["Platform bootstrap", "super_admin", "superadmin", "1"],
            ["Demo tenant admin", "customer_admin", "admin@pharmaglimmora.com", "Admin@123"],
            ["QA Head", "qa_head", "qa@pharmaglimmora.com", "Demo@123"],
            ["Regulatory Affairs", "regulatory_affairs", "ra@pharmaglimmora.com", "Demo@123"],
            ["CSV / Validation Lead", "csv_val_lead", "csv@pharmaglimmora.com", "Demo@123"],
            ["QC Lab Director", "qc_lab_director", "qc@pharmaglimmora.com", "Demo@123"],
            ["IT / CDO", "it_cdo", "it@pharmaglimmora.com", "Demo@123"],
            ["Operations Head", "operations_head", "ops@pharmaglimmora.com", "Demo@123"],
        ],
        col_widths=[1.9, 1.7, 3.0, 1.5])

    add_h2(doc, "1.4  Health smoke-tests (run before walkthrough)")
    add_code(doc,
        "# Frontend\n"
        "curl -i http://localhost:3000/login                            # → 200\n\n"
        "# AI proxy → upstream\n"
        "curl -i http://localhost:3000/api/ai-proxy/api/ai/health       # → 200\n"
        "curl -i http://localhost:3000/api/ai-proxy/api/ai/voice/health # → 200")
    add_body(doc,
        "If the AI proxy returns 502 it usually means the upstream Render free-tier is cold-starting "
        "(30–60 s on first hit). Retry.", italic=True, color=GRAY)

    add_page_break(doc)

    # ─── §2 UI Walkthrough ──────────────────────────────────────────────────
    add_h1(doc, "2. Walkthrough — UI exercises")

    # 2.0 Sign in
    add_h2(doc, "2.0  Sign in as super_admin and create a customer")

    add_step_header(doc, "2.0.1", "Open the login page")
    add_body(doc, "Navigate to http://localhost:3000/login.")
    add_image(doc, SHOTS / "01-login.png", "Screenshot 01 — Login screen")
    add_body(doc, "Inputs:")
    add_bullet(doc, "Work email: superadmin")
    add_bullet(doc, "Passcode: 1")
    add_body(doc, "Click Sign in.")
    add_expected(doc, "Redirects to /admin (Administration Console).")
    add_image(doc, SHOTS / "02-admin-after-login.png", "Screenshot 02 — Administration Console")

    add_step_header(doc, "2.0.2", "Add a new customer tenant")
    add_body(doc, "Click + New Account (top right).")
    add_image(doc, SHOTS / "04-new-account-modal.png", "Screenshot 04 — Add Customer Account modal")
    add_table(doc,
        ["Field", "Value"],
        [
            ["Customer Name", "AI Test Pharma Ltd."],
            ["Username", "aitest_admin"],
            ["Email", "admin@aitest.com"],
            ["New Password", "Test@1234"],
            ["Confirm Password", "Test@1234"],
            ["Language", "English, United States (default)"],
            ["Time Zone", "Asia/Kolkata (default)"],
            ["Require MFA", "off (default)"],
        ],
        col_widths=[2.6, 5.9])
    add_body(doc, "Click Save Account.")
    add_image(doc, SHOTS / "05-new-account-filled.png", "Screenshot 05 — Form filled")
    add_expected(doc, "A subsequent \"Add Subscription Plan\" modal opens automatically.")
    add_image(doc, SHOTS / "06-new-account-created.png", "Screenshot 06 — Subscription Plan prompt")

    add_step_header(doc, "2.0.3", "Attach a subscription so the tenant is usable")
    add_table(doc,
        ["Field", "Value"],
        [
            ["Start date", "today"],
            ["Expiry date", "today + 1 year"],
            ["Max accounts", "15"],
            ["Status", "Yes (Active)"],
        ],
        col_widths=[2.6, 5.9])
    add_body(doc, "Click Save Plan.")
    add_expected(doc, "\"Account and subscription created\" toast.")
    add_image(doc, SHOTS / "08-account-active.png", "Screenshot 08 — Tenant now Active")
    add_body(doc,
        "Note on local date pickers: when the locale is en-IN, the date inputs are interpreted as "
        "DD-MM-YYYY and the new tenant can render as \"No active subscription\" if the parsed dates land "
        "before today. Use admin@pharmaglimmora.com / Admin@123 for the rest of the walkthrough — that "
        "tenant is pre-seeded with an active subscription and exercises identical code paths.",
        italic=True, color=GRAY)

    add_step_header(doc, "2.0.4", "Sign out and re-enter as customer_admin")
    add_body(doc, "Visit /api/auth/signout → Sign out → back to /login.")
    add_body(doc, "Inputs:")
    add_bullet(doc, "Work email: admin@pharmaglimmora.com")
    add_bullet(doc, "Passcode: Admin@123")
    add_expected(doc, "Redirects to / (Dashboard) with the application sidebar populated.")
    add_image(doc, SHOTS / "09-dashboard-as-customer-admin.png", "Screenshot 09 — Dashboard as customer_admin")

    add_page_break(doc)

    # 2.1 AIChatbot
    add_h2(doc, "2.1  AI Chatbot — POST /api/ai/chat   ✅")
    add_body(doc,
        "The floating AI Assistant button is on every authenticated page (bottom right). Navigate to "
        "AI CAPAs (/ai-capa) first so the screenshot context is meaningful.")
    add_image(doc, SHOTS / "10-ai-capa-index.png", "Screenshot 10 — /ai-capa index (empty)")
    add_body(doc, "Click the floating AI Assistant button.")
    add_image(doc, SHOTS / "11-chatbot-opened.png", "Screenshot 11 — Chatbot opened")
    add_body(doc, "Input (text box at bottom of chatbot panel):")
    add_code(doc, "What is a CAPA in GMP?")
    add_body(doc, "Click the Send (paper-plane) icon.")
    add_expected(doc,
        "Browser → POST /api/ai-proxy/api/ai/chat → upstream → 200. Assistant message bubble appears "
        "within 5–20 seconds with a substantive answer. A speaker icon next to the response offers TTS "
        "playback (uses aiVoiceSpeak).")
    add_image(doc, SHOTS / "12-chatbot-response.png", "Screenshot 12 — Chatbot response rendered")

    add_page_break(doc)

    # 2.2 CAPA create
    add_h2(doc, "2.2  CAPA create from AI Generator — POST /api/v1/capa/create   ✅")

    add_step_header(doc, "2.2.1", "Open the modal")
    add_body(doc, "Navigate CAPA Tracker (/capa) → tab \"CAPA Tracker\".")
    add_image(doc, SHOTS / "13-capa-tracker.png", "Screenshot 13 — CAPA module shell")
    add_image(doc, SHOTS / "14-capa-tracker-tab.png", "Screenshot 14 — CAPA Tracker tab")
    add_body(doc, "Click AI CAPA (sparkle icon, right of the \"All sources\" dropdown).")
    add_image(doc, SHOTS / "15-ai-capa-modal.png", "Screenshot 15 — AI CAPA modal")

    add_step_header(doc, "2.2.2", "Fill the form")
    add_table(doc,
        ["Field", "Value"],
        [
            ["Initial severity", "High (default)"],
            ["Problem statement", "Out-of-specification dissolution result detected on Tablet Coater Line 3 batch 26-04-A002. Three of six tablets failed the 30-minute release at the upper RH limit. Initial trend suggests humidity excursion during coating step."],
            ["Source", "Deviation"],
            ["Area affected", "Manufacturing"],
            ["Equipment / Product", "Tablet Coater Line 3 / Batch 26-04-A002"],
            ["Supporting document", "(skip)"],
        ],
        col_widths=[2.0, 6.5])
    add_body(doc,
        "The Problem statement zod schema enforces min(10); shorter inputs raise a client-side validation "
        "error that the AI backend would otherwise reject with an opaque 422.", italic=True, color=GRAY)
    add_image(doc, SHOTS / "16-ai-capa-modal-filled.png", "Screenshot 16 — Form filled")
    add_body(doc, "Click Generate CAPA.")
    add_expected(doc,
        "Browser → POST /api/ai-proxy/api/v1/capa/create (multipart) → 200. Response renders an AI "
        "analysis card with CAPA id (e.g. CAPA-2026-305), AI risk score (0–100 %), pattern detected text, "
        "recurrence alert, AI recommendation, and similar past CAPAs.")
    add_image(doc, SHOTS / "17-capa-created.png", "Screenshot 17 — AI CAPA generated")
    add_body(doc, "Click Accept & close → redirects to /ai-capa/<CAPA-ID> (the lifecycle page).")

    add_page_break(doc)

    # 2.3 Lifecycle stages
    add_h2(doc, "2.3  Lifecycle stages (12 endpoints)   ✅")
    add_body(doc,
        "Single page, same flow per stage. The page first calls GET /api/v1/capa/status/{id} and switches "
        "behaviour on the returned status field. On entry it fans out Promise.allSettled across the five "
        "/<stage>/capa/{id} GET endpoints; stages not yet submitted return 204 via the proxy 404→204 "
        "collapse (silent).")

    add_step_header(doc, "2.3.0", "Initial state — GET /api/v1/capa/status/{id}")
    add_image(doc, SHOTS / "18-ai-capa-lifecycle.png", "Screenshot 18 — Lifecycle page on entry")
    add_body(doc, "Browser network log on entry:")
    add_code(doc, "GET /api/ai-proxy/api/v1/capa/status/CAPA-2026-305 → 200")
    add_body(doc,
        "Status: Open. No by-capa GETs fire while status === \"open\". Submit RCA is the only enabled action.")

    add_step_header(doc, "2.3.1", "Submit RCA — POST /api/v1/rca/submit")
    add_body(doc, "Click Submit RCA.")
    add_image(doc, SHOTS / "19-rca-modal.png", "Screenshot 19 — RCA modal")
    add_table(doc,
        ["Field", "Value"],
        [
            ["RCA method", "5-Why (default)"],
            ["Evidence (optional)", "5-Why: (1) Tablets failed dissolution → (2) Coating film thickness uneven → (3) Coater RH exceeded 65% during spraying → (4) HVAC dehumidifier set point too high → (5) SOP-HVAC-014 last revised 2024, threshold not aligned to current product specs. Root cause: outdated dehumidifier set point in SOP."],
        ],
        col_widths=[2.0, 6.5])
    add_body(doc, "Click Submit.")
    add_body(doc, "Expected network sequence (UI waits ~5–15 seconds for the AI parse):")
    add_code(doc,
        "POST /api/ai-proxy/api/v1/rca/submit                       → 200\n"
        "GET  /api/ai-proxy/api/v1/capa/status/CAPA-2026-305        → 200  (status = \"RCA Submitted\")\n"
        "GET  /api/ai-proxy/api/v1/rca/capa/CAPA-2026-305           → 200  (populated)\n"
        "GET  /api/ai-proxy/api/v1/action-plan/capa/CAPA-2026-305   → 204  silent\n"
        "GET  /api/ai-proxy/api/v1/monitoring/capa/CAPA-2026-305    → 204  silent\n"
        "GET  /api/ai-proxy/api/v1/effectiveness/capa/CAPA-2026-305 → 204  silent\n"
        "GET  /api/ai-proxy/api/v1/closure/capa/CAPA-2026-305       → 204  silent")
    add_body(doc,
        "The RCA card flips to Submitted and shows the parsed rca_id, structured why_1..why_5, root_cause, "
        "and rca_quality_score.")
    add_image(doc, SHOTS / "20-rca-submitted.png", "Screenshot 20 — RCA submitted")

    add_step_header(doc, "2.3.2", "Submit action plan — POST /api/v1/action-plan/submit")
    add_body(doc, "Click Submit action plan.")
    add_image(doc, SHOTS / "21-action-plan-modal.png", "Screenshot 21 — Action plan modal")
    add_table(doc,
        ["Field", "Value"],
        [
            ["Action", "Revise SOP-HVAC-014 to lower dehumidifier set point from 65% RH to 55% RH for tablet coater room"],
            ["Responsible", "Dr. Priya Sharma"],
            ["Due date", "2026-06-30"],
        ],
        col_widths=[2.0, 6.5])
    add_body(doc, "Click Submit.")
    add_expected(doc,
        "POST /api/v1/action-plan/submit → 200. action-plan/capa/{id} flips from 204 → 200.")
    add_image(doc, SHOTS / "22-action-plan-submitted.png", "Screenshot 22 — Action plan submitted")

    add_step_header(doc, "2.3.3", "Submit monitoring check — POST /api/v1/monitoring/check")
    add_body(doc, "Click Submit monitoring check.")
    add_image(doc, SHOTS / "23-monitoring-modal.png", "Screenshot 23 — Monitoring modal")
    add_table(doc,
        ["Field", "Value"],
        [
            ["Action", "Revise SOP-HVAC-014 dehumidifier set point"],
            ["Status", "On Track"],
            ["Note", "Draft revision under QA review. On schedule for 2026-06-30 effective date."],
        ],
        col_widths=[2.0, 6.5])
    add_body(doc,
        "Status must be one of: On Track | In Progress | Overdue | Completed. The older values Delayed | "
        "Blocked are 422-rejected by the backend.", italic=True, color=GRAY)
    add_body(doc, "Click Submit.")
    add_expected(doc, "POST /api/v1/monitoring/check → 200. monitoring/capa/{id} flips from 204 → 200.")
    add_image(doc, SHOTS / "24-monitoring-submitted.png", "Screenshot 24 — Monitoring submitted")

    add_step_header(doc, "2.3.4", "Effectiveness check — POST /api/v1/effectiveness/check")
    add_body(doc, "Click Run effectiveness check.")
    add_image(doc, SHOTS / "25-effectiveness-modal.png", "Screenshot 25 — Effectiveness modal")
    add_table(doc,
        ["Field", "Value"],
        [
            ["Days since CAPA", "90 (default)"],
            ["New issues reported?", "No (default)"],
            ["Evidence", "at least one — defaults are fine"],
        ],
        col_widths=[2.0, 6.5])
    add_body(doc, "Click Submit.")
    add_expected(doc,
        "POST /api/v1/effectiveness/check → 200. Effectiveness card shows effectiveness_score, "
        "effectiveness_rating (HIGHLY_EFFECTIVE | PARTIALLY EFFECTIVE | NEEDS_IMPROVEMENT), and the "
        "capa_can_be_closed flag.")
    add_image(doc, SHOTS / "26-effectiveness-submitted.png", "Screenshot 26 — Effectiveness submitted")

    add_step_header(doc, "2.3.5", "Initiate closure — POST /api/v1/closure/initiate")
    add_body(doc, "Click Initiate closure.")
    add_image(doc, SHOTS / "27-closure-modal.png", "Screenshot 27 — Closure modal")
    add_table(doc,
        ["Field", "Value"],
        [
            ["Approved by", "Dr. Priya Sharma"],
            ["Designation", "QA Head"],
            ["Electronic signature", "PS-SIGN-2026-305"],
            ["Related CAPAs reviewed?", "Yes (default)"],
            ["Document changes approved?", "Yes (default)"],
            ["Closure rationale", "SOP-HVAC-014 revised, effective 2026-06-30. Three follow-up batches passed dissolution at first attempt. Coater RH held under 55% throughout. Effective per ICH Q10 review."],
        ],
        col_widths=[2.5, 6.0])
    add_body(doc, "Click Initiate.")
    add_expected(doc,
        "POST /api/v1/closure/initiate → 200. Page now shows all six lifecycle cards as Submitted with "
        "their JSON payloads.")
    add_image(doc, SHOTS / "28-capa-closed.png", "Screenshot 28 — Lifecycle complete")

    add_step_header(doc, "2.3.6", "Stage-status lookups — GET /api/v1/<stage>/status/{id}")
    add_body(doc, "Navigate to AI Tools (/ai-tools).")
    add_image(doc, SHOTS / "30-ai-tools.png", "Screenshot 30 — AI Tools page")
    add_body(doc,
        "Five identical forms — paste the matching id (RCA-2026-103, AP-2026-…, MON-2026-…, EFF-2026-501, "
        "CLOSURE-2026-… from the lifecycle JSON) into the corresponding box and click Submit. Each returns "
        "the canonical status payload for that stage.")
    add_table(doc,
        ["Endpoint", "URL pattern"],
        [
            ["RCA status", "GET /api/v1/rca/status/{rca_id}"],
            ["Action plan status", "GET /api/v1/action-plan/status/{action_plan_id}"],
            ["Monitoring status", "GET /api/v1/monitoring/status/{monitoring_id}"],
            ["Effectiveness status", "GET /api/v1/effectiveness/status/{effectiveness_id}"],
            ["Closure status", "GET /api/v1/closure/status/{closure_id}"],
        ],
        col_widths=[2.4, 6.1])

    add_page_break(doc)

    # 2.4 CAPA listing & dismiss
    add_h2(doc, "2.4  CAPA listing & alert dismissal")

    add_h3(doc, "2.4.1  GET /api/v1/capa/customer/{customer_id}   ✅")
    add_body(doc,
        "Already exercised on the /ai-capa index page (every navigation refresh). Browser network log:")
    add_code(doc,
        "GET /api/ai-proxy/api/v1/capa/customer/cmp6zc3si0001r0gcbijkf7hc → 200")
    add_body(doc,
        "The customer id is the Prisma tenant id of the logged-in user.")

    add_h3(doc, "2.4.2  GET /api/v1/capa/all   🛠️")
    add_body(doc,
        "Not surfaced in the UI on this branch — the /ai-capa index is scoped per-customer by design. "
        "Verify via curl:")
    add_code(doc,
        "curl -sS http://localhost:3000/api/ai-proxy/api/v1/capa/all")
    add_expected(doc,
        "{\"total\": N, \"capas\": [ {capa_id, problem_statement, source, severity, status, is_recurring, risk_score, created_at}, … ]}")

    add_h3(doc, "2.4.3  POST /api/v1/capa/dismiss-alert   🛠️")
    add_body(doc,
        "Used by the \"Dismiss alert\" link on the AI-Generated CAPA modal and by the recurrence-alert "
        "banner on the /ai-capa index. Verified by hand via curl:")
    add_code(doc,
        "curl -sS -X POST http://localhost:3000/api/ai-proxy/api/v1/capa/dismiss-alert \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -H \"auth: anonymous\" \\\n"
        "  --data-raw '{\n"
        "    \"capa_id\": \"CAPA-2026-305\",\n"
        "    \"alert_type\": \"recurrence\",\n"
        "    \"dismissal_reason\": \"Verified - no similar CAPAs in prior 12 months\",\n"
        "    \"electronic_signature\": \"PS-DISMISS-001\",\n"
        "    \"dismissed_by\": \"Dr. Priya Sharma\"\n"
        "  }'")
    add_expected(doc,
        "200 with an envelope containing the new audit_id and dismissed_at timestamp.")
    add_body(doc,
        "Common gotcha: the FastAPI body parser is strict about Unicode in JSON literals. Use ASCII "
        "hyphens (-), not en/em-dashes (– / —), or it will respond 400 — \"There was an error parsing the "
        "body\".", italic=True, color=GRAY)

    add_page_break(doc)

    # 2.5 Diagnostics
    add_h2(doc, "2.5  Diagnostics & users — health checks + users list   ✅")
    add_body(doc, "On /ai-tools, scroll to Diagnostics. Three Ping buttons:")
    add_table(doc,
        ["Button", "Endpoint"],
        [
            ["AI health", "GET /api/ai/health"],
            ["Voice health", "GET /api/ai/voice/health"],
            ["Users list", "GET /api/v1/users/  (307 → 308 trailing-slash redirect chain, then 200)"],
        ],
        col_widths=[1.5, 7.0])
    add_body(doc, "Click each — each turns green and renders the raw JSON below. Network log:")
    add_code(doc,
        "GET /api/ai-proxy/api/ai/health         → 200\n"
        "GET /api/ai-proxy/api/ai/voice/health   → 200\n"
        "GET /api/ai-proxy/api/v1/users/         → 308 → /api/v1/users → 307 → /api/v1/users/ → 200")

    # 2.6 Audit
    add_h2(doc, "2.6  Audit trail")

    add_h3(doc, "2.6.1  GET /api/v1/audit/all   🛠️")
    add_body(doc,
        "Not surfaced in the UI on this branch (the in-app /audit-trail page reads from local Prisma, "
        "not from the AI backend). Verify via curl through the proxy:")
    add_code(doc,
        "curl -sS http://localhost:3000/api/ai-proxy/api/v1/audit/all")
    add_expected(doc,
        "{\"total\": N, \"audit_logs\": [ {audit_id, action_type, feature_id, record_id, username, "
        "status, timestamp}, … ]}. N is at minimum 6 after running through the walkthrough (create_capa, "
        "submit_rca, submit_action_plan, submit_monitoring, submit_effectiveness, initiate_closure).")

    add_h3(doc, "2.6.2  GET /api/v1/audit/record/{id}   ✅")
    add_body(doc, "On /ai-tools → Audit record card:")
    add_table(doc,
        ["Field", "Value"],
        [["record_id", "CAPA-2026-305 (the CAPA you just walked through)"]],
        col_widths=[2.0, 6.5])
    add_body(doc, "Click Submit.")
    add_expected(doc,
        "GET /api/v1/audit/record/CAPA-2026-305 → 200. The card expands the JSON response scoped to that "
        "record id.")
    add_image(doc, SHOTS / "31-audit-record-response.png", "Screenshot 31 — Audit record response")

    add_page_break(doc)

    # ─── §3 Curl ────────────────────────────────────────────────────────────
    add_h1(doc, "3. Walkthrough — curl-only exercises")
    add_body(doc,
        "Some endpoints can't sensibly be walked via mouse-and-keyboard (auth/signup is normally driven "
        "by the LoginPage in silent mode; voice transcribe/chat need binary audio uploads). Here is the "
        "exact curl set used to validate them — copy/paste-friendly.")

    add_h2(doc, "3.1  Auth — aiSignup + aiLogin   🛠️")
    add_code(doc,
        "# Signup — 201 + access_token\n"
        "curl -sS -X POST http://localhost:3000/api/ai-proxy/api/v1/auth/signup \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -d '{\n"
        "    \"user_id\":     \"USER-test01\",\n"
        "    \"username\":    \"test_qa\",\n"
        "    \"email\":       \"test@aitest.com\",\n"
        "    \"password\":    \"Test@123\",\n"
        "    \"customer_id\": \"CUST_test01\",\n"
        "    \"role\":        \"qa_manager\"\n"
        "  }'\n\n"
        "# Login — 200 + access_token (different from signup-issued token)\n"
        "curl -sS -X POST http://localhost:3000/api/ai-proxy/api/v1/auth/login \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -d '{\"username\":\"test_qa\",\"password\":\"Test@123\"}'")
    add_expected(doc,
        "Both return JSON with access_token, token_type: \"Bearer\", username, customer_id, role, message. "
        "The access token is an HS256 JWT; the upstream is permissive (\"anonymous\" is accepted) so token "
        "absence elsewhere does NOT 401.")

    add_h2(doc, "3.2  Voice — transcribe / speak / chat   🛠️")
    add_code(doc,
        "# 1. Generate a sample audio clip (TTS) so we have something to send back\n"
        "curl -sS -X POST http://localhost:3000/api/ai-proxy/api/ai/voice/speak \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -d '{\"text\":\"Hello from the audit testing manual.\",\"voice\":\"alloy\"}' \\\n"
        "  -o speak.mp3\n"
        "ls -l speak.mp3                            # ~45 KB audio/mpeg\n\n"
        "# 2. Transcribe it back — must return the same text\n"
        "curl -sS -X POST http://localhost:3000/api/ai-proxy/api/ai/voice/transcribe \\\n"
        "  -F \"audio=@./speak.mp3;type=audio/mpeg\"\n"
        "# → {\"text\":\"Hello from the audit testing manual.\",\"customer_id\":\"anonymous\"}\n\n"
        "# 3. Full voice round-trip — audio in, audio out, plus header metadata\n"
        "curl -sS -X POST http://localhost:3000/api/ai-proxy/api/ai/voice/chat \\\n"
        "  -F \"audio=@./speak.mp3;type=audio/mpeg\" \\\n"
        "  -D voicechat-headers.txt \\\n"
        "  -o voicechat-reply.mp3\n"
        "grep -iE 'x-user-text|x-ai-reply|x-intent' voicechat-headers.txt")
    add_body(doc, "Expected response:")
    add_bullet(doc, "speak: 200 audio/mpeg, ~45 KB. Voices: alloy | echo | fable | onyx | nova | shimmer.")
    add_bullet(doc, "transcribe: 200 JSON {text, customer_id}.")
    add_bullet(doc,
        "voice/chat: 200 audio/mpeg reply body plus three response headers — x-user-text, x-ai-reply, "
        "x-intent (URI-decoded by the client). The proxy preserves Access-Control-Expose-Headers because "
        "it's same-origin.")

    add_page_break(doc)

    # ─── §4 Coverage matrix ─────────────────────────────────────────────────
    add_h1(doc, "4. Endpoint coverage matrix")
    rows = [
        ["1",  "POST /api/v1/auth/signup",                    "§3.1 curl",            "🛠️ 201"],
        ["2",  "POST /api/v1/auth/login",                     "§3.1 curl",            "🛠️ 200"],
        ["3",  "POST /api/ai/chat",                           "§2.1 UI",              "✅ 200"],
        ["4",  "GET /api/ai/health",                          "§2.5 UI",              "✅ 200"],
        ["5",  "POST /api/ai/voice/transcribe",               "§3.2 curl",            "🛠️ 200"],
        ["6",  "POST /api/ai/voice/speak",                    "§3.2 curl",            "🛠️ 200"],
        ["7",  "POST /api/ai/voice/chat",                     "§3.2 curl",            "🛠️ 200"],
        ["8",  "GET /api/ai/voice/health",                    "§2.5 UI",              "✅ 200"],
        ["9",  "POST /api/v1/capa/create",                    "§2.2 UI multipart",    "✅ 200"],
        ["10", "GET /api/v1/capa/all",                        "§2.4.2 curl",          "🛠️ 200"],
        ["11", "GET /api/v1/capa/customer/{cid}",             "§2.4.1 UI entry",      "✅ 200"],
        ["12", "GET /api/v1/capa/status/{id}",                "§2.3.0 UI",            "✅ 200"],
        ["13", "POST /api/v1/capa/dismiss-alert",             "§2.4.3 curl",          "🛠️ 200"],
        ["14", "POST /api/v1/rca/submit",                     "§2.3.1 UI",            "✅ 200"],
        ["15", "GET /api/v1/rca/capa/{id}",                   "§2.3.1 UI",            "✅ 200 / 🟡 204"],
        ["16", "GET /api/v1/rca/status/{rca_id}",             "§2.3.6 UI",            "✅ 200"],
        ["17", "POST /api/v1/action-plan/submit",             "§2.3.2 UI",            "✅ 200"],
        ["18", "GET /api/v1/action-plan/capa/{id}",           "§2.3.2 UI",            "✅ 200 / 🟡 204"],
        ["19", "GET /api/v1/action-plan/status/{id}",         "§2.3.6 UI",            "✅ 200"],
        ["20", "POST /api/v1/monitoring/check",               "§2.3.3 UI",            "✅ 200"],
        ["21", "GET /api/v1/monitoring/capa/{id}",            "§2.3.3 UI",            "✅ 200 / 🟡 204"],
        ["22", "GET /api/v1/monitoring/status/{id}",          "§2.3.6 UI",            "✅ 200"],
        ["23", "POST /api/v1/effectiveness/check",            "§2.3.4 UI",            "✅ 200"],
        ["24", "GET /api/v1/effectiveness/capa/{id}",         "§2.3.4 UI",            "✅ 200 / 🟡 204"],
        ["25", "GET /api/v1/effectiveness/status/{id}",       "§2.3.6 UI",            "✅ 200"],
        ["26", "POST /api/v1/closure/initiate",               "§2.3.5 UI",            "✅ 200"],
        ["27", "GET /api/v1/closure/capa/{id}",               "§2.3.5 UI",            "✅ 200 / 🟡 204"],
        ["28", "GET /api/v1/closure/status/{id}",             "§2.3.6 UI",            "✅ 200"],
        ["29", "GET /api/v1/audit/all",                       "§2.6.1 curl",          "🛠️ 200"],
        ["30", "GET /api/v1/audit/record/{id}",               "§2.6.2 UI",            "✅ 200"],
        ["31", "GET /api/v1/users/",                          "§2.5 UI",              "✅ 200 (via redirects)"],
    ]
    add_table(doc, ["#", "Method + Path", "Verified by", "Result"], rows,
              col_widths=[0.6, 4.4, 2.0, 1.5])
    add_body(doc, "31 of 31 endpoints verified.", bold=True, color=GREEN)

    add_page_break(doc)

    # ─── §5 Notes ───────────────────────────────────────────────────────────
    add_h1(doc, "5. Operational notes for testers")
    notes = [
        ("Cold starts.",
         "Render free-tier sleeps after ~15 min of idle. First request after sleep can take 30–60 s. "
         "Don't tighten the fetch timeout — the UI handles the wait."),
        ("SQLite reset on Render redeploy.",
         "Demo CAPAs / users on the AI backend can vanish after a redeploy. Re-seed by re-running §2.2."),
        ("Console-clean discipline.",
         "silentStatuses: [404] + the proxy's 404→204 collapse together mean every by-capa GET that "
         "fires before its stage is recorded is logged as silent rather than a red error. Red lines in "
         "the console are real bugs — investigate."),
        ("Stage gating.",
         "The lifecycle UI gates everything on capa/status/{id}.status (NOT on stage, which lags by one "
         "submit). Direct-navigating to /ai-capa/<id> for an unknown CAPA results in only the Submit RCA "
         "button being shown."),
        ("Monitoring enum.",
         "On Track | In Progress | Overdue | Completed. The old Delayed | Blocked values are 422 errors "
         "today."),
        ("Auth header name.",
         "The AI backend reads auth: <token> (NOT Authorization: Bearer …). The unified client and "
         "proxy both honor this. Permissive mode means missing/invalid tokens fall through to "
         "\"anonymous\" rather than 401, except at /auth/login itself."),
    ]
    for title, body in notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        a = p.add_run(title + " ")
        a.bold = True
        a.font.size = Pt(11)
        a.font.color.rgb = BRAND_AMBER
        b = p.add_run(body)
        b.font.size = Pt(11)
        b.font.color.rgb = BRAND_DARK

    add_h1(doc, "6. Re-running this manual")
    add_code(doc,
        "# 1. Boot both services\n"
        "npm run dev          # web + api in one terminal\n\n"
        "# 2. Walk §2 in a real browser, or replay it with Playwright:\n"
        "#    Screenshots are written to docs/test-screenshots/.\n\n"
        "# 3. Sanity-check the network panel at each step — every line should be either\n"
        "#    200, 201, or 204. A 404 outside the proxy's collapse path is a regression.")
    add_body(doc, "Done correctly the entire walkthrough takes ~12 minutes including the AI-side waits.")

    # Save
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
