"""
Build the Pharma Glimmora user manual as TWO Word documents:

  1. Pharma-Glimmora-AI-Manual.docx       — full walkthrough, every screen
  2. Pharma-Glimmora-AI-Quick-Test.docx   — minimal smoke test (~10 steps)

Run from the project root:
    python docs/manual/build_manual.py

Both files land in docs/manual/. Screenshots come from
docs/manual/screenshots/ (captured separately via Playwright).
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("python-docx is required. Install with: pip install python-docx\n")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
OUT_FULL = ROOT / "Pharma-Glimmora-AI-Manual.docx"
OUT_QUICK = ROOT / "Pharma-Glimmora-AI-Quick-Test.docx"

BRAND_AMBER = RGBColor(0x8B, 0x69, 0x14)
BRAND_DARK = RGBColor(0x30, 0x2D, 0x29)
GRAY = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x10, 0xB9, 0x81)
RED = RGBColor(0xDC, 0x26, 0x26)


# ── Style helpers ───────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_DARK


def add_para(doc, text, *, bold=False, color=None, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_step(doc, n, title, detail=""):
    """Numbered step with bold title and optional detail."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    num = p.add_run(f"{n}. ")
    num.bold = True
    num.font.color.rgb = BRAND_AMBER
    num.font.size = Pt(11)
    t = p.add_run(title)
    t.bold = True
    t.font.size = Pt(11)
    if detail:
        d = p.add_run(f"  {detail}")
        d.font.size = Pt(10)
        d.font.color.rgb = GRAY


def add_check(doc, text):
    """Acceptance criterion bullet."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"✓  {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = GREEN


def add_warn(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"!  {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RED


def add_image(doc, filename, caption=None, width=6.0):
    img_path = SHOTS / filename
    if not img_path.exists():
        add_para(doc, f"[missing screenshot: {filename}]", italic=True, color=GRAY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f"Figure — {caption}")
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = GRAY


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rPr.append(rFonts)


def add_kv_table(doc, rows, header=("Field", "Value")):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = BRAND_AMBER
        set_cell_shading(hdr[i], "FFF7E6")
    for i, (k, v) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = k
        cells[1].text = v
        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)


def title_block(doc, title, subtitle, version):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pharma Glimmora")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BRAND_AMBER

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title)
    r2.font.size = Pt(18)
    r2.font.color.rgb = BRAND_DARK

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.italic = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = GRAY

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(version)
    r4.font.size = Pt(9)
    r4.font.color.rgb = GRAY


# ── Content blocks shared by both documents ────────────────────────────────

LOGIN_QA = ("qa@pharmaglimmora.com", "QaHead@123")


def credentials_block(doc):
    add_heading(doc, "Test credentials", level=2)
    add_kv_table(doc, [
        ("URL", "http://localhost:3000"),
        ("Email", LOGIN_QA[0]),
        ("Passcode", LOGIN_QA[1]),
        ("Role", "QA Head (covers most lifecycle steps)"),
    ])


# ╔═══════════════════════════════════════════════════════════════════════╗
# ║ FULL MANUAL                                                           ║
# ╚═══════════════════════════════════════════════════════════════════════╝

def build_full(doc):
    title_block(doc,
                "AI Backend — Full User Manual",
                "End-to-end walkthrough for manual testing",
                "v3.0 · 2026-05-04")
    doc.add_page_break()

    # ── Overview ───────────────────────────────────────────────────────────
    add_heading(doc, "1. Overview", level=1)
    add_para(doc,
        "The Pharma Glimmora AI backend is integrated into the existing "
        "compliance platform. This manual walks every AI-powered surface "
        "in the order a real QA Head would touch them — sign-in, AI CAPA "
        "creation, the 5-stage CAPA lifecycle, the AI assistant chatbot, "
        "the audit trail, and the AI Tools diagnostics page.")
    add_para(doc,
        "All AI traffic is routed through a same-origin proxy at "
        "/api/ai-proxy so the browser never sees CORS errors when the "
        "Render free-tier dyno cold-starts. There is nothing to configure "
        "— just sign in.")

    credentials_block(doc)
    doc.add_paragraph()

    # ── Sign in ────────────────────────────────────────────────────────────
    add_heading(doc, "2. Sign in", level=1)
    add_step(doc, 1, "Open the app", "Visit http://localhost:3000 — the login page loads.")
    add_step(doc, 2, "Enter the QA credentials",
             "Email qa@pharmaglimmora.com, Passcode QaHead@123, then click Sign in.")
    add_step(doc, 3, "Wait for redirect",
             "On success you land on the Executive Dashboard. AI access token is fetched automatically.")
    add_image(doc, "01-login.png", "Sign-in page")
    add_check(doc, "Top-right shows 'Dr. Priya Sharma · QA Head'.")
    add_check(doc, "No red errors in the browser console (F12 → Console).")
    doc.add_paragraph()

    add_image(doc, "02-dashboard.png", "Executive Dashboard after sign-in")
    doc.add_page_break()

    # ── CAPA Tracker ──────────────────────────────────────────────────────
    add_heading(doc, "3. AI-generated CAPA", level=1)
    add_para(doc,
        "The AI CAPA flow analyses your problem statement against historical "
        "CAPAs, detects recurrence patterns, and proposes a recommendation. "
        "It also flags risk score and (when applicable) a recurrence alert "
        "that requires an electronic-signature dismissal.")

    add_heading(doc, "3.1  Open the CAPA Tracker", level=2)
    add_step(doc, 1, "Navigate to /capa", "Sidebar → Core Compliance → CAPA Tracker.")
    add_step(doc, 2, "Switch to the CAPA Tracker tab", "Top tab strip on the page.")
    add_image(doc, "03-capa-tracker.png", "CAPA Tracker tab — AI CAPA button is in the toolbar")

    add_heading(doc, "3.2  Generate an AI CAPA", level=2)
    add_step(doc, 1, "Click 'AI CAPA'", "Opens the AI-generated CAPA dialog.")
    add_step(doc, 2, "Fill the form", "")
    add_kv_table(doc, [
        ("Initial severity", "High"),
        ("Problem statement", "(min 10 chars) Tablet hardness OOS — 4.2 kP vs 6–10 kP spec on TX-2026-099. Compression force log shows downward drift."),
        ("Source", "Deviation Report"),
        ("Area affected", "Manufacturing - Compression Suite 1"),
        ("Equipment / Product", "Tablet Press (Korsch XL-100)"),
    ])
    add_step(doc, 3, "Click 'Generate CAPA'",
             "Backend takes 5–15 s. You'll see the generated CAPA-XXXX-XXX with a risk score and AI recommendation.")
    add_step(doc, 4, "Click 'Accept & close'", "The CAPA appears in the tracker table and you land on its lifecycle page.")
    add_image(doc, "04-ai-capa-modal.png", "AI CAPA generation modal")
    add_check(doc, "CAPA ID assigned (e.g. CAPA-2026-310).")
    add_check(doc, "Risk score and AI recommendation are populated.")
    add_check(doc, "Page navigates to /ai-capa/CAPA-XXXX-XXX.")
    add_warn(doc, "If you submit the same recurring problem twice you get a recurrence alert with a Dismiss-alert button — that requires a justification ≥ 5 chars and an electronic signature.")
    doc.add_page_break()

    # ── Lifecycle ─────────────────────────────────────────────────────────
    add_heading(doc, "4. CAPA lifecycle (5 stages)", level=1)
    add_para(doc,
        "Each CAPA progresses through 5 backend-tracked stages. The lifecycle "
        "page (/ai-capa/{capa_id}) gates each stage on the previous one, so "
        "you must submit them in order.")
    add_image(doc, "06-ai-capa-lifecycle.png", "Lifecycle page — 5 cards, gated by stage")

    add_heading(doc, "4.1  Submit RCA", level=2)
    add_step(doc, 1, "Click 'Submit RCA'", "Opens the RCA dialog.")
    add_step(doc, 2, "Pick a method", "e.g. Fishbone Diagram (any value works).")
    add_step(doc, 3, "Enter evidence",
             "Free-text root cause description. Example: 'Compression force calibration drifted out of spec due to wear on tablet press feedback sensor.'")
    add_step(doc, 4, "Submit", "Backend returns RCA-XXXX-XXX. Lifecycle status flips to 'RCA Submitted'.")
    add_check(doc, "RCA card now shows the assigned RCA ID and root cause.")
    add_check(doc, "Action Plan card unlocks.")

    add_heading(doc, "4.2  Submit Action Plan", level=2)
    add_step(doc, 1, "Click 'Submit action plan'", "")
    add_step(doc, 2, "Enter at least one action row", "")
    add_kv_table(doc, [
        ("Action", "Recalibrate Korsch XL-100 compression force sensor; reduce maintenance interval."),
        ("Responsible", "Mr. Rajesh Kumar (Maintenance Lead)"),
        ("Due date", "Pick any future date, e.g. 2026-06-15"),
    ])
    add_step(doc, 3, "Submit", "Returns AP-XXXX-XXX. Status flips to 'Action Plan Submitted'.")

    add_heading(doc, "4.3  Submit Monitoring Check", level=2)
    add_step(doc, 1, "Click 'Submit monitoring check'", "")
    add_step(doc, 2, "Status",
             "Pick from On Track / In Progress / Overdue / Completed (these match the backend enum exactly — 'Delayed' / 'Blocked' will fail).")
    add_step(doc, 3, "Note", "Free-text progress note.")
    add_step(doc, 4, "Submit", "Returns MON-XXXX-XXX. Status flips to 'Monitoring Complete' once Completed.")

    add_heading(doc, "4.4  Run Effectiveness Check", level=2)
    add_step(doc, 1, "Click 'Run effectiveness check'", "")
    add_step(doc, 2, "Fill the form",
             "Days since CAPA = 90, New issues = No, Action description, Completed = Yes, Evidence = Attached, Note, plus a Metric block (e.g. Tablet hardness 4.2 → 7.8 kP).")
    add_step(doc, 3, "Submit", "Returns EFF-XXXX-XXX.")
    add_check(doc, "Effectiveness verdict shown (Effective / Partially Effective / Not Effective).")

    add_heading(doc, "4.5  Initiate Closure", level=2)
    add_step(doc, 1, "Click 'Initiate closure'", "")
    add_step(doc, 2, "Fill the closure form", "")
    add_kv_table(doc, [
        ("Approved by", "Dr. Priya Sharma"),
        ("Designation", "QA Head"),
        ("Electronic signature", "QA-PSHARMA-2026-XXX (any unique string)"),
        ("Related CAPAs reviewed?", "Yes"),
        ("Document changes approved?", "Yes"),
        ("Closure rationale", "≥ 10 chars, e.g. '90-day window passed with no recurrence; SOP-MFG-042 updated.'"),
    ])
    add_step(doc, 3, "Submit", "Returns CLO-XXXX-XXX. CAPA status flips to 'Closed'.")
    add_check(doc, "All five lifecycle cards now show data.")
    doc.add_page_break()

    # ── Chatbot ───────────────────────────────────────────────────────────
    add_heading(doc, "5. AI Assistant (chatbot)", level=1)
    add_step(doc, 1, "Click the floating bubble", "Bottom-right of every page.")
    add_step(doc, 2, "Send a text message",
             "Try 'How many CAPAs do we have?' or 'Summarise the latest open deviations.'")
    add_step(doc, 3, "Click the speaker on the assistant reply",
             "Plays back the response via OpenAI TTS.")
    add_step(doc, 4, "Try voice input",
             "Hold the round mic button to record (round-trip with assistant reply played aloud), or the pencil mic for dictation into the input. Voice features need a real microphone and browser permission.")
    add_image(doc, "10-chatbot-open.png", "Chatbot drawer")
    add_check(doc, "Reply streams in within ~5 s.")
    add_check(doc, "Speaker icon plays audio with no console errors.")
    add_warn(doc, "First voice request asks for mic permission — allow it. If denied, only text + speaker work.")
    doc.add_page_break()

    # ── Audit trail ───────────────────────────────────────────────────────
    add_heading(doc, "6. Audit trail", level=1)
    add_step(doc, 1, "Navigate to /audit-trail", "Sidebar → Readiness & Governance → Audit Trail.")
    add_step(doc, 2, "Switch to the 'AI Backend' tab",
             "Pulls /api/v1/audit/all from the backend (every signature event since signup).")
    add_image(doc, "09-audit-ai.png", "AI Backend audit log")
    add_check(doc, "Recent CAPA / RCA / Closure events appear with timestamps and audit IDs.")

    # ── AI Tools ──────────────────────────────────────────────────────────
    add_heading(doc, "7. AI Tools (diagnostic page)", level=1)
    add_para(doc,
        "Direct lookups for every AI backend endpoint by ID. Useful for spot-"
        "checking what was just submitted, or for verifying that a stage's "
        "audit record actually exists.")
    add_step(doc, 1, "Navigate to /ai-tools", "")
    add_step(doc, 2, "Ping the diagnostics", "AI Assistant health, AI Voice health, Users — all should return 200 with JSON.")
    add_step(doc, 3, "Look up a stage by ID",
             "Paste an RCA / Action Plan / Monitoring / Effectiveness / Closure / Audit ID returned earlier and click Submit.")
    add_image(doc, "07-ai-tools.png", "AI Tools page")
    add_check(doc, "All three health pings return 200.")
    add_check(doc, "ID lookups for the records you submitted return their full payload.")
    doc.add_page_break()

    # ── Pass criteria ─────────────────────────────────────────────────────
    add_heading(doc, "8. Pass criteria", level=1)
    add_check(doc, "Sign-in works and aiAccessToken is cached (check via DevTools → Application → Local Storage → glimmora-state).")
    add_check(doc, "AI CAPA creation returns a CAPA ID with risk score and AI recommendation.")
    add_check(doc, "All 5 lifecycle stages submit cleanly and the status field advances correctly.")
    add_check(doc, "Recurrence alert (when triggered) requires reason + signature before dismissal.")
    add_check(doc, "Chatbot replies arrive within ~5 s; speaker playback works.")
    add_check(doc, "Audit trail AI tab is populated with the events you triggered.")
    add_check(doc, "AI Tools health pings + ID lookups all return 200.")
    add_check(doc, "Browser console shows 0 errors at every page (browser-native 4xx logs are fine — those are the only ones you can't suppress and they don't indicate a bug).")

    add_heading(doc, "9. Known caveats", level=1)
    add_warn(doc, "Render free-tier dyno sleeps after ~15 min idle. First request after wake-up takes ~30 s — this is expected, not a bug.")
    add_warn(doc, "Render's SQLite DB is not persistent across cold restarts. CAPAs you created earlier may return 404 the next day; just create fresh ones.")
    add_warn(doc, "Voice round-trip needs HTTPS or localhost AND a real mic. Playwright cannot exercise it.")


# ╔═══════════════════════════════════════════════════════════════════════╗
# ║ QUICK SMOKE TEST                                                      ║
# ╚═══════════════════════════════════════════════════════════════════════╝

def build_quick(doc):
    title_block(doc,
                "AI Backend — Quick Smoke Test",
                "Shortest happy-path manual test (~10 min)",
                "v3.0 · 2026-05-04")
    doc.add_paragraph()

    credentials_block(doc)
    doc.add_paragraph()

    add_heading(doc, "Goal", level=2)
    add_para(doc,
        "Confirm that login + AI CAPA creation + the full 5-stage lifecycle "
        "+ chatbot all work end-to-end against the deployed backend. ~10 "
        "minutes start to finish.")

    add_heading(doc, "Steps", level=2)

    add_step(doc, 1, "Sign in",
             "http://localhost:3000 → email/passcode above → Sign in.")
    add_image(doc, "01-login.png", "Login page", width=4.5)
    add_check(doc, "Top-right shows 'Dr. Priya Sharma · QA Head'.")

    add_step(doc, 2, "Open CAPA Tracker",
             "Sidebar → CAPA Tracker → click the 'CAPA Tracker' tab.")
    add_image(doc, "03-capa-tracker.png", "CAPA Tracker tab", width=4.5)

    add_step(doc, 3, "Generate an AI CAPA",
             "Click 'AI CAPA'. Fill: severity High; problem '≥ 10 chars'; source / area / equipment any short text. Click Generate CAPA.")
    add_image(doc, "04-ai-capa-modal.png", "AI CAPA modal", width=4.5)
    add_check(doc, "After 5–15 s a CAPA-XXXX-XXX appears with risk score + AI recommendation.")

    add_step(doc, 4, "Accept & open lifecycle",
             "Click 'Accept & close'. You land on /ai-capa/CAPA-XXXX-XXX.")

    add_step(doc, 5, "Submit RCA",
             "Click 'Submit RCA'. Method = Fishbone, evidence = any sentence. Submit.")
    add_check(doc, "RCA card populates; Action Plan unlocks.")

    add_step(doc, 6, "Submit Action Plan",
             "Click 'Submit action plan'. Action / Responsible / Due date — any values. Submit.")

    add_step(doc, 7, "Monitoring check",
             "Click 'Submit monitoring check'. Status MUST be one of: On Track / In Progress / Overdue / Completed. Submit.")
    add_warn(doc, "Do not pick anything outside that 4-value enum — backend will return 422.")

    add_step(doc, 8, "Effectiveness",
             "Click 'Run effectiveness check'. Days = 90, New issues = No, Completed = Yes, Evidence = Attached, plus any metric block. Submit.")

    add_step(doc, 9, "Closure",
             "Click 'Initiate closure'. Approved by / Designation / Signature / both Yes / rationale ≥ 10 chars. Submit.")
    add_image(doc, "06-ai-capa-lifecycle.png", "Lifecycle page after all 5 stages", width=4.5)
    add_check(doc, "All 5 cards show submitted data; CAPA status = Closed.")

    add_step(doc, 10, "Chatbot",
             "Click the floating bubble (bottom-right). Type 'How many CAPAs do we have?' and send. Click the speaker icon on the reply.")
    add_image(doc, "10-chatbot-open.png", "Chatbot drawer", width=4.5)
    add_check(doc, "Reply arrives in ~5 s. Speaker plays audio.")

    add_heading(doc, "Pass / Fail", level=2)
    add_check(doc, "Every step above completed without a red error toast in the UI.")
    add_check(doc, "Browser console (F12 → Console) shows 0 JavaScript errors. (Browser-native 'Failed to load resource' 404 lines are not bugs — ignore them.)")
    add_check(doc, "Audit trail at /audit-trail → AI Backend tab shows your events with timestamps.")

    add_warn(doc, "If the very first request hangs ~30 s — that's Render's free-tier waking up. Wait, then retry once.")


# ── Build ──────────────────────────────────────────────────────────────────

def main():
    if not SHOTS.exists():
        sys.stderr.write(f"Screenshot directory missing: {SHOTS}\n")
        sys.exit(1)

    full = Document()
    build_full(full)
    full.save(OUT_FULL)
    print(f"wrote {OUT_FULL.relative_to(ROOT.parent.parent)}")

    quick = Document()
    build_quick(quick)
    quick.save(OUT_QUICK)
    print(f"wrote {OUT_QUICK.relative_to(ROOT.parent.parent)}")


if __name__ == "__main__":
    main()
